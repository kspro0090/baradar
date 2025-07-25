#!/usr/bin/env python3
"""
PDF Generator Module with Persian Font Support
Handles PDF generation from DOCX templates with proper Persian font handling
"""

import os
import re
from typing import Dict, List, Optional, Tuple
from io import BytesIO
import logging

# ReportLab imports
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.pdfmetrics import registerFontFamily
from reportlab.lib.enums import TA_RIGHT, TA_LEFT, TA_CENTER, TA_JUSTIFY

# For RTL text handling
from arabic_reshaper import reshape
from bidi.algorithm import get_display

# For DOCX parsing
from docx import Document
from docx.shared import Pt, RGBColor

# Setup logging
logger = logging.getLogger(__name__)

class PersianPDFGenerator:
    """Handles PDF generation with proper Persian/Arabic font support"""
    
    def __init__(self, fonts_dir: str = "fonts"):
        """
        Initialize the PDF generator with font directory
        
        Args:
            fonts_dir: Directory containing font files
        """
        self.fonts_dir = fonts_dir
        self.registered_fonts = {}
        self.default_font = 'Helvetica'  # Fallback font
        self.persian_font = None
        
        # Register Persian fonts on initialization
        self._register_persian_fonts()
        
    def _register_persian_fonts(self):
        """Register Persian fonts with ReportLab"""
        font_files = {
            'Vazirmatn': 'Vazirmatn-Regular.ttf',
            'Vazirmatn-Bold': 'Vazirmatn-Bold.ttf',
            'Vazirmatn-Light': 'Vazirmatn-Light.ttf',
            'Vazirmatn-Medium': 'Vazirmatn-Medium.ttf'
        }
        
        # Try to register Vazirmatn fonts
        for font_name, font_file in font_files.items():
            font_path = os.path.join(self.fonts_dir, font_file)
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    self.registered_fonts[font_name.lower()] = font_name
                    logger.info(f"Registered font: {font_name} from {font_path}")
                    
                    # Set first successful font as Persian font
                    if not self.persian_font:
                        self.persian_font = font_name
                except Exception as e:
                    logger.warning(f"Failed to register font {font_name}: {str(e)}")
            else:
                logger.debug(f"Font file not found: {font_path}")
        
        # Register font family if we have regular and bold
        if 'vazirmatn' in self.registered_fonts and 'vazirmatn-bold' in self.registered_fonts:
            try:
                registerFontFamily(
                    'Vazirmatn',
                    normal='Vazirmatn',
                    bold='Vazirmatn-Bold',
                    italic='Vazirmatn',  # Use regular for italic
                    boldItalic='Vazirmatn-Bold'  # Use bold for bold italic
                )
                logger.info("Registered Vazirmatn font family")
            except Exception as e:
                logger.warning(f"Failed to register font family: {str(e)}")
    
    def _map_font_name(self, font_name: str) -> str:
        """
        Map font names from documents to registered fonts
        
        Args:
            font_name: Original font name from document
            
        Returns:
            Mapped font name or default font
        """
        if not font_name:
            return self.persian_font or self.default_font
            
        # Normalize font name
        normalized = font_name.lower().strip()
        
        # Direct mapping
        if normalized in self.registered_fonts:
            return self.registered_fonts[normalized]
        
        # Common font mappings
        font_mappings = {
            'vazir': 'Vazirmatn',
            'b nazanin': 'Vazirmatn',
            'bnazanin': 'Vazirmatn',
            'iransans': 'Vazirmatn',
            'iran sans': 'Vazirmatn',
            'tahoma': 'Vazirmatn',
            'arial': self.default_font,
            'times new roman': self.default_font,
            'calibri': self.default_font
        }
        
        # Check mappings
        for key, value in font_mappings.items():
            if key in normalized:
                if value == 'Vazirmatn' and self.persian_font:
                    return self.persian_font
                return value
        
        # If font not found, use Persian font for Persian text, default otherwise
        return self.persian_font or self.default_font
    
    def _process_rtl_text(self, text: str) -> str:
        """
        Process text for proper RTL display
        
        Args:
            text: Original text
            
        Returns:
            Processed text for RTL display
        """
        if not text:
            return text
            
        # Check if text contains Persian/Arabic characters
        if not any('\u0600' <= char <= '\u06FF' or '\u0750' <= char <= '\u077F' for char in text):
            return text
        
        try:
            # Reshape Arabic/Persian text
            reshaped_text = reshape(text)
            # Apply bidirectional algorithm
            bidi_text = get_display(reshaped_text)
            return bidi_text
        except Exception as e:
            logger.warning(f"RTL processing failed: {str(e)}")
            return text
    
    def _create_paragraph_style(self, 
                              font_name: str = None,
                              font_size: int = 12,
                              alignment: str = 'RIGHT',
                              text_color: str = '#000000',
                              bold: bool = False,
                              italic: bool = False) -> ParagraphStyle:
        """
        Create a paragraph style with proper font handling
        
        Args:
            font_name: Font name from document
            font_size: Font size in points
            alignment: Text alignment (LEFT, RIGHT, CENTER, JUSTIFY)
            text_color: Text color in hex format
            bold: Whether text is bold
            italic: Whether text is italic
            
        Returns:
            ParagraphStyle object
        """
        # Map font name
        mapped_font = self._map_font_name(font_name)
        
        # Handle bold variant
        if bold and f"{mapped_font}-Bold" in self.registered_fonts.values():
            mapped_font = f"{mapped_font}-Bold"
        
        # Map alignment
        alignment_map = {
            'LEFT': TA_LEFT,
            'RIGHT': TA_RIGHT,
            'CENTER': TA_CENTER,
            'JUSTIFY': TA_JUSTIFY
        }
        
        style = ParagraphStyle(
            name='CustomStyle',
            fontName=mapped_font,
            fontSize=font_size,
            leading=font_size * 1.2,
            alignment=alignment_map.get(alignment.upper(), TA_RIGHT),
            textColor=text_color,
            wordWrap='RTL' if mapped_font == self.persian_font else 'LTR'
        )
        
        return style
    
    def generate_pdf_from_docx(self, 
                              docx_path: str,
                              output_path: str,
                              replacements: Dict[str, str] = None,
                              page_size=A4) -> bool:
        """
        Generate PDF from DOCX template with placeholder replacements
        
        Args:
            docx_path: Path to DOCX template file
            output_path: Path for output PDF file
            replacements: Dictionary of placeholders to replace
            page_size: Page size (default A4)
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Load DOCX document
            doc = Document(docx_path)
            
            # Create PDF document
            pdf_buffer = BytesIO()
            pdf_doc = SimpleDocTemplate(
                pdf_buffer,
                pagesize=page_size,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Story elements
            story = []
            
            # Process paragraphs
            for para in doc.paragraphs:
                if not para.text.strip():
                    story.append(Spacer(1, 0.2 * inch))
                    continue
                
                # Get paragraph text
                para_text = para.text
                
                # Replace placeholders
                if replacements:
                    for placeholder, value in replacements.items():
                        para_text = para_text.replace(placeholder, str(value))
                
                # Process RTL text
                para_text = self._process_rtl_text(para_text)
                
                # Get paragraph formatting
                if para.runs:
                    # Use first run's formatting
                    run = para.runs[0]
                    font_name = run.font.name
                    font_size = run.font.size.pt if run.font.size else 12
                    bold = run.font.bold if run.font.bold is not None else False
                    italic = run.font.italic if run.font.italic is not None else False
                    
                    # Get color
                    color = '#000000'
                    if run.font.color and run.font.color.rgb:
                        rgb = run.font.color.rgb
                        color = f'#{rgb.red:02x}{rgb.green:02x}{rgb.blue:02x}'
                else:
                    font_name = None
                    font_size = 12
                    bold = False
                    italic = False
                    color = '#000000'
                
                # Determine alignment
                alignment = 'RIGHT'  # Default for Persian
                if para.alignment:
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    if para.alignment == WD_ALIGN_PARAGRAPH.LEFT:
                        alignment = 'LEFT'
                    elif para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        alignment = 'CENTER'
                    elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        alignment = 'JUSTIFY'
                
                # Create paragraph style
                style = self._create_paragraph_style(
                    font_name=font_name,
                    font_size=font_size,
                    alignment=alignment,
                    text_color=color,
                    bold=bold,
                    italic=italic
                )
                
                # Add paragraph to story
                p = Paragraph(para_text, style)
                story.append(p)
                story.append(Spacer(1, 0.1 * inch))
            
            # Process tables if any
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text
                        
                        # Replace placeholders
                        if replacements:
                            for placeholder, value in replacements.items():
                                cell_text = cell_text.replace(placeholder, str(value))
                        
                        # Process RTL text
                        cell_text = self._process_rtl_text(cell_text)
                        row_data.append(cell_text)
                    
                    table_data.append(row_data)
                
                if table_data:
                    # Create table with RTL support
                    t = Table(table_data)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
                        ('FONTNAME', (0, 0), (-1, -1), self.persian_font or self.default_font),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(t)
                    story.append(Spacer(1, 0.2 * inch))
            
            # Build PDF
            pdf_doc.build(story)
            
            # Save to file
            pdf_buffer.seek(0)
            with open(output_path, 'wb') as f:
                f.write(pdf_buffer.read())
            
            logger.info(f"PDF generated successfully: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            return False
    
    def generate_pdf_from_text(self,
                              content: str,
                              output_path: str,
                              title: str = None,
                              font_name: str = None,
                              font_size: int = 12,
                              page_size=A4) -> bool:
        """
        Generate PDF from plain text content
        
        Args:
            content: Text content
            output_path: Path for output PDF file
            title: Optional title
            font_name: Font to use
            font_size: Font size
            page_size: Page size
            
        Returns:
            True if successful, False otherwise
        """
        try:
            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                output_path,
                pagesize=page_size,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )
            
            # Story elements
            story = []
            
            # Add title if provided
            if title:
                title_style = self._create_paragraph_style(
                    font_name=font_name,
                    font_size=font_size + 4,
                    alignment='CENTER',
                    bold=True
                )
                title_text = self._process_rtl_text(title)
                story.append(Paragraph(title_text, title_style))
                story.append(Spacer(1, 0.3 * inch))
            
            # Create content style
            content_style = self._create_paragraph_style(
                font_name=font_name,
                font_size=font_size,
                alignment='RIGHT'
            )
            
            # Process content
            paragraphs = content.split('\n\n')
            for para in paragraphs:
                if para.strip():
                    para_text = self._process_rtl_text(para)
                    story.append(Paragraph(para_text, content_style))
                    story.append(Spacer(1, 0.2 * inch))
            
            # Build PDF
            pdf_doc.build(story)
            
            logger.info(f"PDF generated successfully: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            return False


# Convenience function for backward compatibility
def generate_pdf_from_docx(docx_path: str,
                          output_path: str,
                          replacements: Dict[str, str] = None,
                          fonts_dir: str = "fonts") -> bool:
    """
    Generate PDF from DOCX with Persian font support
    
    Args:
        docx_path: Path to DOCX template
        output_path: Path for output PDF
        replacements: Placeholder replacements
        fonts_dir: Directory containing fonts
        
    Returns:
        True if successful, False otherwise
    """
    generator = PersianPDFGenerator(fonts_dir=fonts_dir)
    return generator.generate_pdf_from_docx(docx_path, output_path, replacements)


if __name__ == "__main__":
    # Test the PDF generator
    import tempfile
    
    # Create test DOCX
    test_doc = Document()
    test_doc.add_heading('تست سند فارسی', 0)
    test_doc.add_paragraph('این یک متن تست برای {{employee_name}} است.')
    test_doc.add_paragraph('تاریخ: {{date}}')
    
    # Save test DOCX
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        test_doc.save(tmp.name)
        docx_path = tmp.name
    
    # Test replacements
    replacements = {
        '{{employee_name}}': 'علی احمدی',
        '{{date}}': '1402/10/15'
    }
    
    # Generate PDF
    generator = PersianPDFGenerator()
    success = generator.generate_pdf_from_docx(
        docx_path,
        'test_output.pdf',
        replacements
    )
    
    if success:
        print("PDF generated successfully!")
    else:
        print("PDF generation failed!")
    
    # Cleanup
    os.unlink(docx_path)