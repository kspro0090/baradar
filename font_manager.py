"""
Font Management Module
Handles font detection, validation, and registration for document processing
"""

import os
import re
from typing import List, Dict, Set, Tuple, Optional
from docx import Document
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from werkzeug.utils import secure_filename
import logging

logger = logging.getLogger(__name__)


class FontManager:
    """Manages fonts for document processing and PDF generation"""
    
    def __init__(self, font_directory: str = 'static/fonts'):
        self.font_directory = font_directory
        self.registered_fonts = set()
        self._ensure_font_directory()
        self._scan_and_register_fonts()
    
    def _ensure_font_directory(self):
        """Ensure the font directory exists"""
        os.makedirs(self.font_directory, exist_ok=True)
    
    def _scan_and_register_fonts(self):
        """Scan font directory and register all fonts with ReportLab"""
        if not os.path.exists(self.font_directory):
            return
        
        for filename in os.listdir(self.font_directory):
            if filename.lower().endswith(('.ttf', '.otf')):
                font_path = os.path.join(self.font_directory, filename)
                font_name = self._get_font_name_from_file(filename)
                
                try:
                    # Register with ReportLab
                    pdfmetrics.registerFont(TTFont(font_name, font_path))
                    self.registered_fonts.add(font_name)
                    logger.info(f"Registered font: {font_name} from {filename}")
                except Exception as e:
                    logger.error(f"Failed to register font {filename}: {str(e)}")
    
    def _get_font_name_from_file(self, filename: str) -> str:
        """Extract font name from filename"""
        # Remove extension
        name = os.path.splitext(filename)[0]
        # Remove common suffixes
        name = re.sub(r'[-_](Regular|Bold|Italic|Light|Medium|Heavy|Black|Thin)$', '', name, flags=re.IGNORECASE)
        return name
    
    def extract_fonts_from_docx(self, docx_path: str) -> Dict[str, Set[str]]:
        """
        Extract all fonts used in a DOCX file
        Returns a dictionary with font usage information
        """
        fonts_used = set()
        font_details = {
            'all_fonts': set(),
            'paragraph_fonts': set(),
            'table_fonts': set(),
            'header_footer_fonts': set(),
            'default_font': None
        }
        
        try:
            doc = Document(docx_path)
            
            # Get default font from document styles
            if doc.styles.default('Normal').font and doc.styles.default('Normal').font.name:
                default_font = doc.styles.default('Normal').font.name
                font_details['default_font'] = default_font
                fonts_used.add(default_font)
            
            # Extract fonts from paragraphs
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    if run.font and run.font.name:
                        fonts_used.add(run.font.name)
                        font_details['paragraph_fonts'].add(run.font.name)
            
            # Extract fonts from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.font and run.font.name:
                                    fonts_used.add(run.font.name)
                                    font_details['table_fonts'].add(run.font.name)
            
            # Extract fonts from headers and footers
            for section in doc.sections:
                # Header
                for paragraph in section.header.paragraphs:
                    for run in paragraph.runs:
                        if run.font and run.font.name:
                            fonts_used.add(run.font.name)
                            font_details['header_footer_fonts'].add(run.font.name)
                
                # Footer
                for paragraph in section.footer.paragraphs:
                    for run in paragraph.runs:
                        if run.font and run.font.name:
                            fonts_used.add(run.font.name)
                            font_details['header_footer_fonts'].add(run.font.name)
            
            font_details['all_fonts'] = fonts_used
            
        except Exception as e:
            logger.error(f"Error extracting fonts from DOCX: {str(e)}")
            raise
        
        return font_details
    
    def is_font_available(self, font_name: str) -> Tuple[bool, Optional[str]]:
        """
        Check if a font is available in the system
        Returns (is_available, matched_font_file)
        """
        if not font_name:
            return False, None
        
        # Normalize font name for comparison
        normalized_font = self._normalize_font_name(font_name)
        
        # Check registered fonts first
        for registered in self.registered_fonts:
            if self._normalize_font_name(registered) == normalized_font:
                return True, registered
        
        # Check font files in directory
        if os.path.exists(self.font_directory):
            for filename in os.listdir(self.font_directory):
                if filename.lower().endswith(('.ttf', '.otf')):
                    file_font_name = self._get_font_name_from_file(filename)
                    if self._normalize_font_name(file_font_name) == normalized_font:
                        return True, filename
        
        return False, None
    
    def _normalize_font_name(self, font_name: str) -> str:
        """Normalize font name for comparison (case-insensitive, remove spaces/hyphens)"""
        if not font_name:
            return ""
        
        # Convert to lowercase
        normalized = font_name.lower()
        
        # Remove common font suffixes
        suffixes = ['regular', 'bold', 'italic', 'light', 'medium', 'heavy', 'black', 'thin']
        for suffix in suffixes:
            normalized = normalized.replace(f'-{suffix}', '').replace(f'_{suffix}', '').replace(f' {suffix}', '')
        
        # Remove spaces, hyphens, underscores
        normalized = re.sub(r'[\s\-_]+', '', normalized)
        
        return normalized
    
    def check_missing_fonts(self, font_list: Set[str]) -> Dict[str, bool]:
        """
        Check which fonts from a list are missing
        Returns a dictionary: {font_name: is_available}
        """
        font_status = {}
        
        for font in font_list:
            is_available, _ = self.is_font_available(font)
            font_status[font] = is_available
        
        return font_status
    
    def get_missing_fonts(self, font_list: Set[str]) -> List[str]:
        """Get list of missing fonts"""
        font_status = self.check_missing_fonts(font_list)
        return [font for font, available in font_status.items() if not available]
    
    def save_uploaded_font(self, file, original_filename: str) -> Tuple[bool, str]:
        """
        Save an uploaded font file
        Returns (success, message)
        """
        try:
            # Validate file extension
            if not original_filename.lower().endswith(('.ttf', '.otf')):
                return False, "Invalid font file format. Only TTF and OTF files are supported."
            
            # Secure the filename
            filename = secure_filename(original_filename)
            if not filename:
                filename = f"font_{os.urandom(8).hex()}.ttf"
            
            # Save the file
            filepath = os.path.join(self.font_directory, filename)
            file.save(filepath)
            
            # Try to register the font
            font_name = self._get_font_name_from_file(filename)
            try:
                pdfmetrics.registerFont(TTFont(font_name, filepath))
                self.registered_fonts.add(font_name)
                return True, f"Font '{font_name}' uploaded and registered successfully."
            except Exception as e:
                # Remove the file if registration fails
                os.unlink(filepath)
                return False, f"Failed to register font: {str(e)}"
                
        except Exception as e:
            logger.error(f"Error saving font file: {str(e)}")
            return False, f"Error saving font file: {str(e)}"
    
    def get_available_fonts(self) -> List[Dict[str, str]]:
        """Get list of all available fonts with details"""
        fonts = []
        
        if os.path.exists(self.font_directory):
            for filename in os.listdir(self.font_directory):
                if filename.lower().endswith(('.ttf', '.otf')):
                    filepath = os.path.join(self.font_directory, filename)
                    font_name = self._get_font_name_from_file(filename)
                    
                    fonts.append({
                        'name': font_name,
                        'filename': filename,
                        'path': filepath,
                        'size': os.path.getsize(filepath),
                        'registered': font_name in self.registered_fonts
                    })
        
        return sorted(fonts, key=lambda x: x['name'].lower())
    
    def delete_font(self, filename: str) -> Tuple[bool, str]:
        """
        Delete a font file
        Returns (success, message)
        """
        try:
            filepath = os.path.join(self.font_directory, filename)
            
            if not os.path.exists(filepath):
                return False, "Font file not found."
            
            # Remove from registered fonts
            font_name = self._get_font_name_from_file(filename)
            self.registered_fonts.discard(font_name)
            
            # Delete the file
            os.unlink(filepath)
            
            return True, f"Font '{font_name}' deleted successfully."
            
        except Exception as e:
            logger.error(f"Error deleting font: {str(e)}")
            return False, f"Error deleting font: {str(e)}"
    
    def get_font_for_pdf(self, requested_font: str) -> str:
        """
        Get the best available font for PDF generation
        Falls back to default fonts if requested font is not available
        """
        # Check if requested font is available
        is_available, matched_font = self.is_font_available(requested_font)
        if is_available:
            return self._get_font_name_from_file(matched_font) if matched_font else requested_font
        
        # Fallback fonts for Persian/Arabic text
        persian_fallbacks = ['Vazir', 'IRANSans', 'B Nazanin', 'B Mitra']
        for fallback in persian_fallbacks:
            is_available, matched_font = self.is_font_available(fallback)
            if is_available:
                logger.warning(f"Font '{requested_font}' not found, using '{fallback}' as fallback")
                return self._get_font_name_from_file(matched_font) if matched_font else fallback
        
        # Final fallback
        logger.warning(f"Font '{requested_font}' not found, using system default")
        return 'Helvetica'  # ReportLab's default font


# Global font manager instance
font_manager = FontManager()