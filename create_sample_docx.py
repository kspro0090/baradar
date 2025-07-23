#!/usr/bin/env python3
"""
Create a sample Word document template for testing the preview feature
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_sample_template():
    """Create a sample Word template with placeholders"""
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('فرم درخواست خدمات', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a paragraph with some Persian text
    p1 = doc.add_paragraph('این یک نمونه قالب برای درخواست خدمات است. در این قالب از متغیرهایی استفاده شده که با اطلاعات فرم جایگزین خواهند شد.')
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add personal information section
    doc.add_heading('اطلاعات شخصی', level=1)
    
    p2 = doc.add_paragraph('نام و نام خانوادگی: ')
    p2.add_run('{{FULLNAME}}').bold = True
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p3 = doc.add_paragraph('کد ملی: ')
    p3.add_run('{{NATIONALCODE}}').bold = True
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p4 = doc.add_paragraph('شماره تماس: ')
    p4.add_run('{{PHONE}}').bold = True
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add address section
    doc.add_heading('آدرس', level=1)
    
    p5 = doc.add_paragraph('آدرس کامل: ')
    p5.add_run('{{ADDRESS}}').bold = True
    p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add description section
    doc.add_heading('توضیحات', level=1)
    
    p6 = doc.add_paragraph('{{DESCRIPTION}}')
    p6.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add footer
    doc.add_paragraph()
    footer = doc.add_paragraph('این سند به صورت خودکار تولید شده است.')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save the document
    uploads_dir = 'uploads/docx_templates'
    os.makedirs(uploads_dir, exist_ok=True)
    
    filename = os.path.join(uploads_dir, 'sample_template.docx')
    doc.save(filename)
    
    print(f"Sample template created: {filename}")
    print("\nPlaceholders used:")
    print("- {{FULLNAME}} - نام و نام خانوادگی")
    print("- {{NATIONALCODE}} - کد ملی")
    print("- {{PHONE}} - شماره تماس")
    print("- {{ADDRESS}} - آدرس")
    print("- {{DESCRIPTION}} - توضیحات")

if __name__ == '__main__':
    create_sample_template()