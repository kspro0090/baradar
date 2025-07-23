"""
Document Processing Service
Handles local document processing, placeholder replacement, and PDF generation
"""

import os
import io
import re
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Image
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_RIGHT, TA_CENTER, TA_LEFT
import arabic_reshaper
from bidi.algorithm import get_display
from bs4 import BeautifulSoup
import base64


class DocumentProcessor:
    """Process documents locally without modifying the original"""
    
    def __init__(self, font_manager=None):
        # Use provided font manager or create a new one
        if font_manager:
            self.font_manager = font_manager
        else:
            from font_manager import FontManager
            self.font_manager = FontManager()
        
        # Register fonts is now handled by font manager
        # No need to register fonts here
    
    def process_docx_template(self, docx_data, replacements):
        """Process DOCX template and replace placeholders while preserving formatting"""
        # Create a temporary file to work with
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(docx_data)
            tmp_file_path = tmp_file.name
        
        try:
            # Open the document
            doc = Document(tmp_file_path)
            
            # First, detect all placeholders in the document
            placeholders_found = self._detect_placeholders_in_docx(doc)
            print(f"Found placeholders in document: {placeholders_found}")
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                self._replace_placeholders_in_paragraph(paragraph, replacements)
            
            # Replace placeholders in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_placeholders_in_paragraph(paragraph, replacements)
            
            # Replace placeholders in headers and footers
            for section in doc.sections:
                # Header
                for paragraph in section.header.paragraphs:
                    self._replace_placeholders_in_paragraph(paragraph, replacements)
                # Footer
                for paragraph in section.footer.paragraphs:
                    self._replace_placeholders_in_paragraph(paragraph, replacements)
            
            # Save the modified document
            output_path = tmp_file_path.replace('.docx', '_filled.docx')
            doc.save(output_path)
            
            # Read the filled document
            with open(output_path, 'rb') as f:
                filled_data = f.read()
            
            # Clean up
            os.unlink(output_path)
            
            return filled_data
            
        finally:
            # Clean up temp file
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)
    
    def _detect_placeholders_in_docx(self, doc):
        """Detect all placeholders in the document"""
        placeholders = set()
        
        # Check paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text
            found = re.findall(r'\{\{([^}]+)\}\}', text)
            placeholders.update(found)
        
        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        found = re.findall(r'\{\{([^}]+)\}\}', text)
                        placeholders.update(found)
        
        # Check headers and footers
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                text = paragraph.text
                found = re.findall(r'\{\{([^}]+)\}\}', text)
                placeholders.update(found)
            for paragraph in section.footer.paragraphs:
                text = paragraph.text
                found = re.findall(r'\{\{([^}]+)\}\}', text)
                placeholders.update(found)
        
        return placeholders
    
    def _replace_placeholders_in_paragraph(self, paragraph, replacements):
        """Replace placeholders in a paragraph while preserving formatting"""
        # Get the full text of the paragraph
        full_text = paragraph.text
        
        # Check if any placeholders exist in this paragraph
        has_placeholders = False
        for placeholder, value in replacements.items():
            placeholder_pattern = f'{{{{{placeholder}}}}}'
            if placeholder_pattern in full_text:
                has_placeholders = True
                break
        
        if not has_placeholders:
            return
        
        # If paragraph has placeholders, we need to handle runs carefully
        # First, collect all runs and their properties
        runs_data = []
        for run in paragraph.runs:
            runs_data.append({
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name if run.font else None,
                'font_size': run.font.size if run.font else None,
                'font_color': run.font.color.rgb if run.font and run.font.color and run.font.color.rgb else None,
                'highlight_color': run.font.highlight_color if run.font else None
            })
        
        # Reconstruct the full text from runs
        full_text = ''.join(run['text'] for run in runs_data)
        
        # Replace all placeholders in the full text
        for placeholder, value in replacements.items():
            placeholder_pattern = f'{{{{{placeholder}}}}}'
            full_text = full_text.replace(placeholder_pattern, value)
        
        # Clear the paragraph
        paragraph.clear()
        
        # If we have formatting information, try to preserve it
        if runs_data and len(runs_data) > 0:
            # Use the formatting from the first run as default
            default_format = runs_data[0]
            
            # Add the replaced text with the default formatting
            run = paragraph.add_run(full_text)
            if default_format['bold'] is not None:
                run.bold = default_format['bold']
            if default_format['italic'] is not None:
                run.italic = default_format['italic']
            if default_format['underline'] is not None:
                run.underline = default_format['underline']
            if default_format['font_name'] and run.font:
                run.font.name = default_format['font_name']
            if default_format['font_size'] and run.font:
                run.font.size = default_format['font_size']
            if default_format['font_color'] and run.font and run.font.color:
                run.font.color.rgb = default_format['font_color']
            if default_format['highlight_color'] and run.font:
                run.font.highlight_color = default_format['highlight_color']
        else:
            # No formatting info, just add the text
            paragraph.add_run(full_text)
    
    def process_html_template(self, html_data, replacements):
        """Process HTML template and replace placeholders"""
        html_content = html_data.decode('utf-8')
        
        # Replace placeholders
        for placeholder, value in replacements.items():
            placeholder_pattern = f'{{{{{placeholder}}}}}'
            html_content = html_content.replace(placeholder_pattern, value)
        
        return html_content.encode('utf-8')
    
    def _prepare_persian_text(self, text):
        """Prepare Persian text for proper RTL display in PDF"""
        if not text:
            return ""
        # Reshape Arabic/Persian text
        reshaped_text = arabic_reshaper.reshape(text)
        # Apply RTL algorithm
        bidi_text = get_display(reshaped_text)
        return bidi_text
    
    def generate_pdf_from_html(self, html_data, output_path):
        """Generate PDF from HTML content with Persian support"""
        try:
            # Parse HTML
            soup = BeautifulSoup(html_data, 'html.parser')
            
            # Create PDF document
            doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Container for the 'Flowable' objects
            elements = []
            
            # Define styles
            styles = getSampleStyleSheet()
            
            # Persian style
            persian_style = ParagraphStyle(
                'Persian',
                parent=styles['Normal'],
                fontName='Vazir',
                fontSize=12,
                leading=20,
                alignment=TA_RIGHT,
                wordWrap='RTL',
                direction='RTL'
            )
            
            # Process HTML content
            for element in soup.body.children if soup.body else []:
                if element.name == 'p':
                    text = element.get_text(strip=True)
                    if text:
                        # Prepare Persian text
                        persian_text = self._prepare_persian_text(text)
                        para = Paragraph(persian_text, persian_style)
                        elements.append(para)
                        elements.append(Spacer(1, 12))
                
                elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                    text = element.get_text(strip=True)
                    if text:
                        # Create heading style
                        heading_style = ParagraphStyle(
                            f'Heading{element.name[1]}',
                            parent=persian_style,
                            fontSize=20 - int(element.name[1]) * 2,
                            fontName='Vazir-Bold',
                            spaceAfter=12
                        )
                        persian_text = self._prepare_persian_text(text)
                        para = Paragraph(persian_text, heading_style)
                        elements.append(para)
                        elements.append(Spacer(1, 12))
                
                elif element.name == 'table':
                    # Process tables
                    table_data = []
                    for row in element.find_all('tr'):
                        row_data = []
                        for cell in row.find_all(['td', 'th']):
                            text = cell.get_text(strip=True)
                            persian_text = self._prepare_persian_text(text)
                            row_data.append(Paragraph(persian_text, persian_style))
                        if row_data:
                            table_data.append(row_data)
                    
                    if table_data:
                        t = Table(table_data)
                        t.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
                            ('FONTNAME', (0, 0), (-1, -1), 'Vazir'),
                            ('FONTSIZE', (0, 0), (-1, 0), 12),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)
                        ]))
                        elements.append(t)
                        elements.append(Spacer(1, 12))
            
            # Build PDF
            doc.build(elements)
            return True
            
        except Exception as e:
            print(f"Error generating PDF from HTML: {str(e)}")
            return False
    
    def generate_pdf_from_docx(self, docx_data, output_path):
        """Generate PDF from DOCX content with Persian support and layout preservation"""
        # Create a temporary file for the DOCX
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(docx_data)
            tmp_file_path = tmp_file.name
        
        try:
            # Open the document
            doc = Document(tmp_file_path)
            
            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                output_path,
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=18
            )
            
            # Container for the 'Flowable' objects
            elements = []
            
            # Define styles
            styles = getSampleStyleSheet()
            
            # Persian style
            persian_style = ParagraphStyle(
                'Persian',
                parent=styles['Normal'],
                fontName=self.font_manager.get_font_for_pdf('Vazir'),
                fontSize=12,
                leading=20,
                alignment=TA_RIGHT,
                wordWrap='RTL',
                direction='RTL'
            )
            
            # Persian heading style
            persian_heading = ParagraphStyle(
                'PersianHeading',
                parent=persian_style,
                fontSize=16,
                fontName=self.font_manager.get_font_for_pdf('Vazir-Bold'),
                spaceAfter=12,
                spaceBefore=12
            )
            
            # Process document elements
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    # Process paragraph
                    para = None
                    for p in doc.paragraphs:
                        if p._element == element:
                            para = p
                            break
                    
                    if para and para.text.strip():
                        text = para.text.strip()
                        
                        # Check if paragraph has specific font
                        para_font = None
                        if para.runs and para.runs[0].font and para.runs[0].font.name:
                            para_font = self.font_manager.get_font_for_pdf(para.runs[0].font.name)
                        
                        # Check if it's a heading based on style
                        if para.style and para.style.name and para.style.name.startswith('Heading'):
                            # Prepare Persian text for heading
                            persian_text = self._prepare_persian_text(text)
                            
                            # Create custom heading style with document font
                            custom_heading = ParagraphStyle(
                                'CustomHeading',
                                parent=persian_heading
                            )
                            if para_font:
                                custom_heading.fontName = para_font
                            
                            pdf_para = Paragraph(persian_text, custom_heading)
                        else:
                            # Regular paragraph
                            persian_text = self._prepare_persian_text(text)
                            
                            # Create custom style based on paragraph formatting
                            para_style = ParagraphStyle(
                                'CustomPersian',
                                parent=persian_style
                            )
                            
                            # Use paragraph's font if available
                            if para_font:
                                para_style.fontName = para_font
                            
                            # Check alignment
                            if para.alignment:
                                if str(para.alignment) == 'CENTER':
                                    para_style.alignment = TA_CENTER
                                elif str(para.alignment) == 'LEFT':
                                    para_style.alignment = TA_LEFT
                            
                            pdf_para = Paragraph(persian_text, para_style)
                        
                        elements.append(pdf_para)
                        elements.append(Spacer(1, 6))
                
                elif element.tag.endswith('tbl'):
                    # Process table
                    table = None
                    for t in doc.tables:
                        if t._element == element:
                            table = t
                            break
                    
                    if table:
                        table_data = []
                        for row in table.rows:
                            row_data = []
                            for cell in row.cells:
                                cell_text = cell.text.strip()
                                persian_text = self._prepare_persian_text(cell_text)
                                
                                # Create cell paragraph with proper style
                                cell_para = Paragraph(persian_text, persian_style)
                                row_data.append(cell_para)
                            
                            if row_data:
                                table_data.append(row_data)
                        
                        if table_data:
                            # Calculate column widths based on content
                            num_cols = len(table_data[0]) if table_data else 0
                            col_width = (pdf_doc.width) / num_cols if num_cols > 0 else pdf_doc.width
                            
                            t = Table(table_data, colWidths=[col_width] * num_cols)
                            
                            # Apply table style
                            table_style = TableStyle([
                                ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
                                ('FONTNAME', (0, 0), (-1, -1), self.font_manager.get_font_for_pdf('Vazir')),
                                ('FONTSIZE', (0, 0), (-1, -1), 10),
                                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                                ('TOPPADDING', (0, 0), (-1, -1), 8),
                                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                                ('DIRECTION', (0, 0), (-1, -1), 'RTL'),
                            ])
                            
                            # Check if first row should be header
                            if len(table_data) > 1:
                                table_style.add('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey)
                                table_style.add('FONTNAME', (0, 0), (-1, 0), self.font_manager.get_font_for_pdf('Vazir-Bold'))
                            
                            t.setStyle(table_style)
                            elements.append(t)
                            elements.append(Spacer(1, 12))
            
            # Build PDF
            pdf_doc.build(elements)
            return True
            
        except Exception as e:
            print(f"Error generating PDF from DOCX: {str(e)}")
            import traceback
            traceback.print_exc()
            return False
            
        finally:
            # Clean up temp file
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)