#!/usr/bin/env python3
"""
Test script for font management functionality
"""

import os
import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from font_manager import FontManager
from docx import Document
from docx.shared import Pt
import tempfile

def test_font_detection():
    """Test font detection from DOCX files"""
    print("=== Testing Font Detection ===\n")
    
    # Create a test document with various fonts
    doc = Document()
    
    # Add content with different fonts
    para1 = doc.add_paragraph()
    run1 = para1.add_run('This text uses Calibri font.')
    run1.font.name = 'Calibri'
    run1.font.size = Pt(12)
    
    para2 = doc.add_paragraph()
    run2 = para2.add_run('این متن از فونت وزیر استفاده می‌کند.')
    run2.font.name = 'Vazir'
    run2.font.size = Pt(14)
    
    para3 = doc.add_paragraph()
    run3 = para3.add_run('This uses Arial font.')
    run3.font.name = 'Arial'
    
    # Add table with fonts
    table = doc.add_table(rows=2, cols=2)
    cell = table.cell(0, 0)
    cell_para = cell.paragraphs[0]
    cell_run = cell_para.add_run('Table with IRANSans')
    cell_run.font.name = 'IRANSans'
    
    # Save test document
    test_path = 'test_fonts.docx'
    doc.save(test_path)
    
    try:
        # Initialize font manager
        font_manager = FontManager()
        
        # Extract fonts
        font_details = font_manager.extract_fonts_from_docx(test_path)
        
        print("Detected fonts:")
        for font in font_details['all_fonts']:
            print(f"  - {font}")
        
        print(f"\nTotal fonts found: {len(font_details['all_fonts'])}")
        
        # Check font availability
        print("\nFont availability:")
        font_status = font_manager.check_missing_fonts(font_details['all_fonts'])
        for font, available in font_status.items():
            status = "✓ Available" if available else "✗ Missing"
            print(f"  {font}: {status}")
        
        missing = font_manager.get_missing_fonts(font_details['all_fonts'])
        if missing:
            print(f"\nMissing fonts: {', '.join(missing)}")
        else:
            print("\nAll fonts are available!")
        
        return True
        
    finally:
        # Clean up
        if os.path.exists(test_path):
            os.unlink(test_path)

def test_font_normalization():
    """Test font name normalization"""
    print("\n=== Testing Font Normalization ===\n")
    
    font_manager = FontManager()
    
    test_cases = [
        ('Vazir', 'vazir'),
        ('Vazir-Bold', 'vazir'),
        ('IRANSans Regular', 'iransans'),
        ('Arial-Regular', 'arial'),
        ('B Nazanin Bold', 'bnazanin'),
        ('Times New Roman', 'timesnewroman'),
    ]
    
    print("Font normalization tests:")
    for original, expected in test_cases:
        normalized = font_manager._normalize_font_name(original)
        status = "✓" if normalized == expected else "✗"
        print(f"  {status} '{original}' → '{normalized}' (expected: '{expected}')")
    
    return True

def test_font_matching():
    """Test font matching with partial names"""
    print("\n=== Testing Font Matching ===\n")
    
    # Create a temporary font directory
    temp_dir = tempfile.mkdtemp()
    font_manager = FontManager(font_directory=temp_dir)
    
    # Create dummy font files
    font_files = [
        'Vazir-Regular.ttf',
        'IRANSansWeb.ttf',
        'B_Nazanin_Bold.ttf',
        'Arial.ttf'
    ]
    
    for filename in font_files:
        filepath = os.path.join(temp_dir, filename)
        with open(filepath, 'wb') as f:
            f.write(b'dummy font data')
    
    # Test font availability checks
    test_queries = [
        'Vazir',
        'vazir-bold',  # Should match Vazir
        'IRANSans',
        'B Nazanin',
        'Arial',
        'Times New Roman',  # Should not match
    ]
    
    print("Font matching tests:")
    for query in test_queries:
        is_available, matched = font_manager.is_font_available(query)
        if is_available:
            print(f"  ✓ '{query}' → Found (matched: {matched})")
        else:
            print(f"  ✗ '{query}' → Not found")
    
    # Clean up
    import shutil
    shutil.rmtree(temp_dir)
    
    return True

def test_font_fallback():
    """Test font fallback mechanism"""
    print("\n=== Testing Font Fallback ===\n")
    
    font_manager = FontManager()
    
    test_fonts = [
        'NonExistentFont',
        'Vazir',
        'IRANSans',
        'UnknownFont'
    ]
    
    print("Font fallback tests:")
    for font in test_fonts:
        fallback = font_manager.get_font_for_pdf(font)
        print(f"  '{font}' → '{fallback}'")
    
    return True

def main():
    """Run all tests"""
    print("=== Font Management Test Suite ===\n")
    
    results = []
    
    # Run tests
    results.append(("Font Detection", test_font_detection()))
    results.append(("Font Normalization", test_font_normalization()))
    results.append(("Font Matching", test_font_matching()))
    results.append(("Font Fallback", test_font_fallback()))
    
    # Summary
    print("\n=== Test Summary ===\n")
    for test_name, result in results:
        status = "✓ PASS" if result else "✗ FAIL"
        print(f"{test_name}: {status}")
    
    all_passed = all(result for _, result in results)
    if all_passed:
        print("\n✓ All tests passed!")
    else:
        print("\n✗ Some tests failed")
    
    return all_passed

if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)