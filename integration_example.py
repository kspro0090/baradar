#!/usr/bin/env python3
"""
Integration Example - How to use the new PDF generator in existing code
"""

import os
from document_processor import (
    generate_pdf_from_docx_template,
    process_service_request_to_pdf,
    validate_font_availability,
    generate_pdf_from_request
)

# Example 1: Direct PDF generation from DOCX template
def example_direct_pdf_generation():
    """Example of direct PDF generation"""
    print("Example 1: Direct PDF Generation")
    print("-" * 40)
    
    # Template and output paths
    template_path = "templates/employee_form.docx"
    output_path = "outputs/employee_report.pdf"
    
    # Data to replace in template
    replacements = {
        "{{employee_name}}": "علی احمدی",
        "{{employee_id}}": "EMP-12345",
        "{{department}}": "فناوری اطلاعات",
        "{{hire_date}}": "1402/05/15",
        "{{salary}}": "50,000,000 ریال"
    }
    
    # Generate PDF
    success = generate_pdf_from_docx_template(
        template_path,
        output_path,
        replacements
    )
    
    if success:
        print(f"✓ PDF generated: {output_path}")
    else:
        print("✗ PDF generation failed")
    
    return success

# Example 2: Mock service request object
class MockServiceRequest:
    """Mock service request for testing"""
    def __init__(self):
        self.tracking_code = "REQ-2024-001"
        self.created_at = __import__('datetime').datetime.now()
        self.user = type('User', (), {'username': 'admin'})()
        self.service = type('Service', (), {
            'docx_template_path': 'templates/service_template.docx'
        })()
        self._form_data = {
            'name': 'محمد رضایی',
            'phone': '09123456789',
            'email': 'mohammad@example.com',
            'request_type': 'درخواست مرخصی',
            'description': 'درخواست مرخصی استحقاقی'
        }
    
    def get_form_data(self):
        return self._form_data

# Example 3: Process service request
def example_process_service_request():
    """Example of processing a service request to PDF"""
    print("\nExample 2: Process Service Request")
    print("-" * 40)
    
    # Create mock service request
    service_request = MockServiceRequest()
    
    # Process to PDF
    output_path = process_service_request_to_pdf(
        service_request,
        service_request.service.docx_template_path
    )
    
    if output_path:
        print(f"✓ PDF generated: {output_path}")
        print(f"  Tracking code: {service_request.tracking_code}")
    else:
        print("✗ PDF generation failed")
    
    return output_path

# Example 4: Check font availability
def example_check_fonts():
    """Example of checking font availability"""
    print("\nExample 3: Check Font Availability")
    print("-" * 40)
    
    # List of fonts to check
    fonts_to_check = [
        'Vazir',
        'B Nazanin',
        'IRANSans',
        'Arial',
        'Times New Roman',
        'Tahoma'
    ]
    
    # Check availability
    availability = validate_font_availability(fonts_to_check)
    
    print("Font Availability:")
    for font, is_available in availability.items():
        status = "✓ Available" if is_available else "✗ Not available"
        print(f"  {font}: {status}")
    
    return availability

# Example 5: Integration with existing app.py code
def example_app_integration():
    """Example showing how to modify existing app.py code"""
    print("\nExample 4: Integration with app.py")
    print("-" * 40)
    
    print("Replace the existing generate_pdf_from_request function with:")
    print("""
# In app.py, replace:
# from app import generate_pdf_from_request

# With:
from document_processor import generate_pdf_from_request

# Or if you want to use ReportLab instead of Google Docs:
def generate_pdf_from_request(service_request):
    '''Generate PDF using ReportLab with Persian font support'''
    from document_processor import process_service_request_to_pdf
    
    # Get template path
    template_path = service_request.service.docx_template_path
    if not template_path:
        raise Exception("No DOCX template configured for this service")
    
    # Generate PDF
    output_path = process_service_request_to_pdf(
        service_request,
        template_path
    )
    
    if output_path:
        return os.path.basename(output_path)
    else:
        raise Exception("Failed to generate PDF")
""")

# Example 6: Create sample templates
def create_sample_templates():
    """Create sample DOCX templates for testing"""
    print("\nCreating Sample Templates")
    print("-" * 40)
    
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    # Ensure directories exist
    os.makedirs("templates", exist_ok=True)
    os.makedirs("outputs", exist_ok=True)
    
    # Template 1: Employee Form
    doc1 = Document()
    doc1.add_heading('فرم اطلاعات کارمند', 0)
    
    p1 = doc1.add_paragraph('نام و نام خانوادگی: {{employee_name}}')
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p2 = doc1.add_paragraph('کد پرسنلی: {{employee_id}}')
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p3 = doc1.add_paragraph('بخش: {{department}}')
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p4 = doc1.add_paragraph('تاریخ استخدام: {{hire_date}}')
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p5 = doc1.add_paragraph('حقوق: {{salary}}')
    p5.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc1.save('templates/employee_form.docx')
    print("✓ Created: templates/employee_form.docx")
    
    # Template 2: Service Request
    doc2 = Document()
    doc2.add_heading('فرم درخواست خدمات', 0)
    
    doc2.add_paragraph('کد پیگیری: {{tracking_code}}')
    doc2.add_paragraph('تاریخ درخواست: {{request_date}}')
    doc2.add_paragraph('درخواست‌دهنده: {{requester_name}}')
    doc2.add_paragraph('')
    doc2.add_paragraph('نام: {{name}}')
    doc2.add_paragraph('تلفن: {{phone}}')
    doc2.add_paragraph('ایمیل: {{email}}')
    doc2.add_paragraph('نوع درخواست: {{request_type}}')
    doc2.add_paragraph('شرح درخواست: {{description}}')
    
    doc2.save('templates/service_template.docx')
    print("✓ Created: templates/service_template.docx")

def main():
    """Run all examples"""
    print("PDF Generator Integration Examples")
    print("=" * 50)
    
    # Create sample templates
    create_sample_templates()
    
    # Run examples
    example_direct_pdf_generation()
    example_process_service_request()
    example_check_fonts()
    example_app_integration()
    
    print("\n" + "=" * 50)
    print("Integration examples completed!")
    
    # List generated files
    if os.path.exists("outputs"):
        files = os.listdir("outputs")
        if files:
            print("\nGenerated PDF files:")
            for f in files:
                if f.endswith('.pdf'):
                    size = os.path.getsize(os.path.join("outputs", f))
                    print(f"  - {f} ({size} bytes)")

if __name__ == "__main__":
    main()