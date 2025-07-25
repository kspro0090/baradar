#!/usr/bin/env python3
"""
Test script for PDF generation with Persian fonts
"""

import os
import sys
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import our PDF generator
from pdf_generator import PersianPDFGenerator

def create_test_docx(filename="test_template.docx"):
    """Create a test DOCX file with Persian content and placeholders"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('گزارش کارمند', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add content with different fonts
    p1 = doc.add_paragraph()
    p1.add_run('نام کارمند: ').bold = True
    p1.add_run('{{employee_name}}')
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p2 = doc.add_paragraph()
    p2.add_run('کد پرسنلی: ').bold = True
    p2.add_run('{{employee_id}}')
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    p3 = doc.add_paragraph()
    p3.add_run('تاریخ استخدام: ').bold = True
    p3.add_run('{{hire_date}}')
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add a paragraph with Vazir font (will be mapped)
    p4 = doc.add_paragraph('این متن با فونت Vazir نوشته شده است.')
    if p4.runs:
        p4.runs[0].font.name = 'Vazir'
        p4.runs[0].font.size = Pt(14)
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'ردیف'
    header_cells[1].text = 'شرح'
    header_cells[2].text = 'مبلغ'
    
    # Data rows
    row1_cells = table.rows[1].cells
    row1_cells[0].text = '1'
    row1_cells[1].text = 'حقوق پایه'
    row1_cells[2].text = '{{base_salary}}'
    
    row2_cells = table.rows[2].cells
    row2_cells[0].text = '2'
    row2_cells[1].text = 'اضافه کاری'
    row2_cells[2].text = '{{overtime_pay}}'
    
    # Save document
    doc.save(filename)
    print(f"Test DOCX created: {filename}")
    return filename

def test_pdf_generation():
    """Test PDF generation with various scenarios"""
    print("Testing Persian PDF Generation")
    print("=" * 50)
    
    # Create test DOCX
    docx_file = create_test_docx()
    
    # Test replacements
    replacements = {
        '{{employee_name}}': 'علی محمدی',
        '{{employee_id}}': '12345',
        '{{hire_date}}': '1401/05/15',
        '{{base_salary}}': '50,000,000 ریال',
        '{{overtime_pay}}': '5,000,000 ریال'
    }
    
    # Initialize PDF generator
    generator = PersianPDFGenerator()
    
    # Test 1: Generate PDF with replacements
    print("\nTest 1: Generating PDF with placeholder replacements...")
    output_file = "test_output_persian.pdf"
    success = generator.generate_pdf_from_docx(
        docx_file,
        output_file,
        replacements
    )
    
    if success:
        print(f"✓ PDF generated successfully: {output_file}")
    else:
        print("✗ PDF generation failed!")
    
    # Test 2: Generate PDF from text
    print("\nTest 2: Generating PDF from plain text...")
    persian_text = """
این یک متن تست فارسی است که برای بررسی عملکرد سیستم تولید PDF نوشته شده است.

سیستم باید بتواند:
- متن فارسی را به درستی نمایش دهد
- از فونت Vazirmatn استفاده کند
- جهت متن را از راست به چپ تنظیم کند
- کاراکترهای فارسی را به شکل متصل نمایش دهد

با تشکر از توجه شما.
"""
    
    text_output = "test_text_persian.pdf"
    success2 = generator.generate_pdf_from_text(
        persian_text,
        text_output,
        title="تست متن فارسی",
        font_name="Vazir"
    )
    
    if success2:
        print(f"✓ Text PDF generated successfully: {text_output}")
    else:
        print("✗ Text PDF generation failed!")
    
    # Test 3: Test font mapping
    print("\nTest 3: Testing font mapping...")
    test_fonts = ['vazir', 'Vazir', 'B Nazanin', 'IRANSans', 'Arial', 'Unknown Font']
    
    for font in test_fonts:
        mapped = generator._map_font_name(font)
        print(f"  {font} → {mapped}")
    
    # Cleanup
    if os.path.exists(docx_file):
        os.remove(docx_file)
        print(f"\nCleaned up: {docx_file}")
    
    print("\n" + "=" * 50)
    print("Testing completed!")
    
    # Check if output files exist
    if os.path.exists(output_file):
        print(f"\nOutput file size: {os.path.getsize(output_file)} bytes")
    if os.path.exists(text_output):
        print(f"Text output file size: {os.path.getsize(text_output)} bytes")

if __name__ == "__main__":
    # Check if fonts directory exists
    if not os.path.exists("fonts"):
        print("Error: 'fonts' directory not found!")
        print("Please create a 'fonts' directory and add Vazirmatn-Regular.ttf")
        sys.exit(1)
    
    if not os.path.exists("fonts/Vazirmatn-Regular.ttf"):
        print("Error: Vazirmatn-Regular.ttf not found in fonts directory!")
        print("Please add the font file to the fonts directory")
        sys.exit(1)
    
    # Run tests
    test_pdf_generation()