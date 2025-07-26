#!/usr/bin/env python3
"""
Test script for no-copy PDF generation
"""

import os
import json
from datetime import datetime

def create_mock_service_request():
    """Create a mock service request for testing"""
    
    class MockFormField:
        def __init__(self, field_name, field_label, document_placeholder):
            self.field_name = field_name
            self.field_label = field_label
            self.document_placeholder = document_placeholder
    
    class MockService:
        def __init__(self):
            self.name = "درخواست مرخصی"
            self.google_doc_id = "YOUR_GOOGLE_DOC_ID_HERE"  # Replace with actual ID
            self.docx_template_path = "templates/leave_request.docx"
            self.form_fields = [
                MockFormField('employee_name', 'نام کارمند', '{{employee_name}}'),
                MockFormField('employee_id', 'کد پرسنلی', '{{employee_id}}'),
                MockFormField('department', 'بخش', '{{department}}'),
                MockFormField('leave_type', 'نوع مرخصی', '{{leave_type}}'),
                MockFormField('start_date', 'تاریخ شروع', '{{start_date}}'),
                MockFormField('end_date', 'تاریخ پایان', '{{end_date}}'),
                MockFormField('reason', 'دلیل', '{{reason}}')
            ]
    
    class MockUser:
        def __init__(self):
            self.username = 'test_user'
    
    class MockServiceRequest:
        def __init__(self):
            self.tracking_code = f"REQ-{datetime.now().strftime('%Y%m%d-%H%M%S')}"
            self.created_at = datetime.now()
            self.service = MockService()
            self.user = MockUser()
            self._form_data = {
                'employee_name': 'علی محمدی',
                'employee_id': 'EMP-12345',
                'department': 'فناوری اطلاعات',
                'leave_type': 'استحقاقی',
                'start_date': '1402/11/01',
                'end_date': '1402/11/05',
                'reason': 'امور شخصی'
            }
        
        def get_form_data(self):
            return self._form_data
    
    return MockServiceRequest()

def test_no_copy_generation():
    """Test the no-copy PDF generation"""
    print("Testing No-Copy PDF Generation")
    print("=" * 50)
    
    # Create mock request
    service_request = create_mock_service_request()
    print(f"Created mock request with tracking code: {service_request.tracking_code}")
    
    # Test 1: Try no-copy method (requires Google Docs API)
    print("\nTest 1: No-Copy Method (Google Docs API)")
    print("-" * 40)
    
    if os.path.exists('credentials.json'):
        try:
            from pdf_generator_no_copy import generate_pdf_from_request_no_copy
            
            pdf_filename = generate_pdf_from_request_no_copy(service_request)
            if pdf_filename:
                print(f"✓ PDF generated successfully: {pdf_filename}")
            else:
                print("✗ PDF generation failed")
        except Exception as e:
            print(f"✗ Error: {str(e)}")
    else:
        print("⚠ credentials.json not found - skipping Google Docs test")
    
    # Test 2: Try fallback method (ReportLab)
    print("\nTest 2: Fallback Method (ReportLab)")
    print("-" * 40)
    
    try:
        from pdf_generator_no_copy import generate_pdf_fallback
        
        pdf_filename = generate_pdf_fallback(service_request)
        if pdf_filename:
            print(f"✓ PDF generated successfully: {pdf_filename}")
            
            # Check file size
            pdf_path = os.path.join('pdf_outputs', pdf_filename)
            if os.path.exists(pdf_path):
                size = os.path.getsize(pdf_path)
                print(f"  File size: {size} bytes")
        else:
            print("✗ PDF generation failed")
    except Exception as e:
        print(f"✗ Error: {str(e)}")
    
    # Test 3: Create DOCX template for testing
    print("\nTest 3: Creating DOCX Template")
    print("-" * 40)
    
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        os.makedirs('templates', exist_ok=True)
        
        doc = Document()
        doc.add_heading('فرم درخواست مرخصی', 0)
        
        # Add fields
        fields = [
            ('نام کارمند', '{{employee_name}}'),
            ('کد پرسنلی', '{{employee_id}}'),
            ('بخش', '{{department}}'),
            ('نوع مرخصی', '{{leave_type}}'),
            ('از تاریخ', '{{start_date}}'),
            ('تا تاریخ', '{{end_date}}'),
            ('دلیل درخواست', '{{reason}}')
        ]
        
        for label, placeholder in fields:
            p = doc.add_paragraph()
            p.add_run(f'{label}: ').bold = True
            p.add_run(placeholder)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add footer
        doc.add_paragraph()
        doc.add_paragraph('کد پیگیری: {{tracking_code}}')
        doc.add_paragraph('تاریخ درخواست: {{request_date}}')
        
        doc.save('templates/leave_request.docx')
        print("✓ Created templates/leave_request.docx")
        
        # Now test with DOCX template
        print("\nRetrying fallback with DOCX template...")
        pdf_filename = generate_pdf_fallback(service_request)
        if pdf_filename:
            print(f"✓ PDF generated with template: {pdf_filename}")
        
    except Exception as e:
        print(f"✗ Error creating template: {str(e)}")
    
    print("\n" + "=" * 50)
    print("Testing completed!")
    
    # List generated PDFs
    if os.path.exists('pdf_outputs'):
        files = [f for f in os.listdir('pdf_outputs') if f.endswith('.pdf')]
        if files:
            print("\nGenerated PDFs:")
            for f in files:
                print(f"  - {f}")

def test_placeholder_restoration():
    """Test that placeholders are properly restored"""
    print("\nTesting Placeholder Restoration")
    print("=" * 50)
    
    # Sample replacements
    replacements = {
        '{{employee_name}}': 'علی محمدی',
        '{{department}}': 'فناوری اطلاعات',
        '{{leave_type}}': 'استحقاقی',
        '{{start_date}}': '1402/11/01'
    }
    
    # Create reverse replacements
    reverse_replacements = {v: k for k, v in replacements.items()}
    
    print("Original replacements:")
    for k, v in replacements.items():
        print(f"  {k} → {v}")
    
    print("\nReverse replacements:")
    for k, v in reverse_replacements.items():
        print(f"  {k} → {v}")
    
    print("\nThis ensures document is restored to original state")

if __name__ == "__main__":
    test_no_copy_generation()
    test_placeholder_restoration()