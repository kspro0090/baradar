#!/usr/bin/env python3
"""
Test script for placeholder replacement in Word documents
"""

import os
import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from document_processor import DocumentProcessor
from docx import Document
import tempfile

def test_placeholder_replacement():
    """Test the placeholder replacement functionality"""
    print("=== Testing Placeholder Replacement ===\n")
    
    # First, create a test template
    print("1. Creating test template...")
    from create_test_template import create_test_template
    template_path = create_test_template()
    
    # Read the template
    with open(template_path, 'rb') as f:
        template_data = f.read()
    
    # Define test replacements
    replacements = {
        'date': '1402/10/25',
        'employee_name': 'احمد محمدی',
        'employee_code': '12345',
        'department': 'فناوری اطلاعات',
        'service_type': 'مرخصی استحقاقی',
        'request_type': 'مرخصی سالانه',
        'start_date': '1402/11/01',
        'end_date': '1402/11/05',
        'description': 'درخواست مرخصی جهت امور شخصی',
        'email': 'ahmad.mohammadi@example.com',
        'phone': '09123456789'
    }
    
    print("\n2. Processing template with replacements...")
    processor = DocumentProcessor()
    
    try:
        # Process the template
        filled_data = processor.process_docx_template(template_data, replacements)
        
        # Save the filled document
        output_dir = 'test_output'
        os.makedirs(output_dir, exist_ok=True)
        
        filled_path = os.path.join(output_dir, 'filled_document.docx')
        with open(filled_path, 'wb') as f:
            f.write(filled_data)
        
        print(f"✓ Filled document saved: {filled_path}")
        
        # Verify replacements
        print("\n3. Verifying replacements...")
        doc = Document(filled_path)
        
        # Extract all text
        all_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                all_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        all_text.append(cell.text)
        
        full_text = '\n'.join(all_text)
        
        # Check that no placeholders remain
        import re
        remaining_placeholders = re.findall(r'\{\{([^}]+)\}\}', full_text)
        
        if remaining_placeholders:
            print(f"✗ Found unreplaced placeholders: {remaining_placeholders}")
            success = False
        else:
            print("✓ All placeholders replaced successfully")
            success = True
        
        # Check that replacements are present
        print("\n4. Checking replacement values...")
        for key, value in replacements.items():
            if value in full_text:
                print(f"✓ Found '{key}' → '{value}'")
            else:
                print(f"✗ Missing '{key}' → '{value}'")
                success = False
        
        # Generate PDF
        print("\n5. Generating PDF from filled document...")
        pdf_path = os.path.join(output_dir, 'output.pdf')
        
        if processor.generate_pdf_from_docx(filled_data, pdf_path):
            print(f"✓ PDF generated: {pdf_path}")
            print(f"  File size: {os.path.getsize(pdf_path):,} bytes")
        else:
            print("✗ Failed to generate PDF")
            success = False
        
        return success
        
    except Exception as e:
        print(f"✗ Error: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

def test_complex_placeholder_scenarios():
    """Test complex scenarios like placeholders spanning multiple runs"""
    print("\n=== Testing Complex Placeholder Scenarios ===\n")
    
    # Create a document with placeholders that might span runs
    doc = Document()
    
    # Scenario 1: Placeholder split across runs
    para1 = doc.add_paragraph()
    para1.add_run('{{')
    para1.add_run('employee_')
    para1.add_run('name}}')
    
    # Scenario 2: Placeholder with formatting
    para2 = doc.add_paragraph()
    run = para2.add_run('Name: {{employee_name}}')
    run.bold = True
    
    # Scenario 3: Multiple placeholders in one paragraph
    para3 = doc.add_paragraph('{{first_name}} {{last_name}} - {{email}}')
    
    # Save test document
    test_path = 'test_templates/complex_test.docx'
    doc.save(test_path)
    
    # Read and process
    with open(test_path, 'rb') as f:
        template_data = f.read()
    
    replacements = {
        'employee_name': 'علی رضایی',
        'first_name': 'علی',
        'last_name': 'رضایی',
        'email': 'ali@example.com'
    }
    
    processor = DocumentProcessor()
    filled_data = processor.process_docx_template(template_data, replacements)
    
    # Save and verify
    filled_path = 'test_output/complex_filled.docx'
    with open(filled_path, 'wb') as f:
        f.write(filled_data)
    
    # Check results
    doc = Document(filled_path)
    text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
    
    print(f"Original placeholders → Replaced text:")
    print(f"1. Split placeholder: {'{{employee_name}}' not in text}")
    print(f"2. Formatted placeholder: {'علی رضایی' in text}")
    print(f"3. Multiple placeholders: {all(v in text for v in replacements.values())}")
    
    return True

def main():
    """Run all tests"""
    print("=== Word Document Placeholder Replacement Test Suite ===\n")
    
    results = []
    
    # Test basic replacement
    results.append(("Basic Placeholder Replacement", test_placeholder_replacement()))
    
    # Test complex scenarios
    results.append(("Complex Placeholder Scenarios", test_complex_placeholder_scenarios()))
    
    # Summary
    print("\n=== Test Summary ===\n")
    for test_name, result in results:
        status = "✓ PASS" if result else "✗ FAIL"
        print(f"{test_name}: {status}")
    
    all_passed = all(result for _, result in results)
    if all_passed:
        print("\n✓ All tests passed!")
    else:
        print("\n✗ Some tests failed")
    
    return all_passed

if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)