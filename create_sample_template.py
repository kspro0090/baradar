from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Create a sample Word template
doc = Document()

# Add title
title = doc.add_heading('درخواست خدمت', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add a paragraph with placeholders
doc.add_paragraph()
p1 = doc.add_paragraph('نام و نام خانوادگی: ')
p1.add_run('{{NAME}}').bold = True

p2 = doc.add_paragraph('شماره تماس: ')
p2.add_run('{{PHONE}}').bold = True

p3 = doc.add_paragraph('آدرس: ')
p3.add_run('{{ADDRESS}}').bold = True

p4 = doc.add_paragraph('توضیحات: ')
p4.add_run('{{DESCRIPTION}}').bold = True

doc.add_paragraph()
doc.add_paragraph('تاریخ: {{DATE}}')

# Save the template
os.makedirs('uploads/docx_templates', exist_ok=True)
doc.save('uploads/docx_templates/sample_template.docx')
print("Sample template created successfully!")