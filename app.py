from flask import Flask, render_template, redirect, url_for, flash, request, send_file, jsonify
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from werkzeug.utils import secure_filename
from functools import wraps
import os
import secrets
import json
from datetime import datetime
import re
import tempfile

from google_docs_service import GoogleDocsService
from document_processor import DocumentProcessor
from wtforms.validators import DataRequired, Email
from wtforms import StringField, IntegerField, TextAreaField, SelectField

from config import Config
from models import db, User, Service, FormField, ServiceRequest
from forms import (LoginForm, CreateAdminForm, ServiceForm, FormFieldForm, 
                   ServiceRequestForm, ApprovalForm, TrackingForm)

app = Flask(__name__)
app.config.from_object(Config)

# Initialize extensions
db.init_app(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'
login_manager.login_message = 'لطفاً وارد شوید.'

# Create directories if they don't exist
os.makedirs(app.config['PDF_OUTPUT_FOLDER'], exist_ok=True)

# Initialize Google Docs service
google_docs_service = None
try:
    google_docs_service = GoogleDocsService(os.path.join(os.path.dirname(__file__), 'credentials.json'))
except Exception as e:
    app.logger.error(f"Failed to initialize Google Docs service: {str(e)}")

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Decorators for role checking
def system_manager_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_system_manager():
            flash('شما دسترسی به این بخش ندارید.', 'danger')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function

def approval_admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_approval_admin():
            flash('شما دسترسی به این بخش ندارید.', 'danger')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function

# Helper functions
def generate_tracking_code():
    """Generate a unique tracking code"""
    while True:
        code = secrets.token_hex(5).upper()
        if not ServiceRequest.query.filter_by(tracking_code=code).first():
            return code



# Routes
@app.route('/')
def index():
    """Public homepage showing available services"""
    services = Service.query.filter_by(is_active=True).all()
    return render_template('user/index.html', services=services)

@app.route('/login', methods=['GET', 'POST'])
def login():
    """Login page for system managers and approval admins"""
    if current_user.is_authenticated:
        if current_user.is_system_manager():
            return redirect(url_for('admin_dashboard'))
        elif current_user.is_approval_admin():
            return redirect(url_for('approver_dashboard'))
    
    form = LoginForm()
    if form.validate_on_submit():
        user = User.query.filter_by(username=form.username.data).first()
        if user and user.check_password(form.password.data):
            login_user(user, remember=form.remember_me.data)
            next_page = request.args.get('next')
            if not next_page:
                if user.is_system_manager():
                    next_page = url_for('admin_dashboard')
                else:
                    next_page = url_for('approver_dashboard')
            return redirect(next_page)
        flash('نام کاربری یا رمز عبور اشتباه است.', 'danger')
    return render_template('login.html', form=form)

@app.route('/logout')
@login_required
def logout():
    logout_user()
    flash('با موفقیت خارج شدید.', 'success')
    return redirect(url_for('index'))

# System Manager Routes
@app.route('/admin')
@login_required
@system_manager_required
def admin_dashboard():
    """System manager dashboard"""
    services = Service.query.all()
    admins = User.query.filter_by(role='approval_admin').all()
    return render_template('admin/dashboard.html', 
                         services=services, 
                         admins=admins)

@app.route('/admin/create-admin', methods=['GET', 'POST'])
@login_required
@system_manager_required
def create_admin():
    """Create new approval admin"""
    form = CreateAdminForm()
    if form.validate_on_submit():
        admin = User(
            username=form.username.data,
            email=form.email.data,
            role='approval_admin'
        )
        admin.set_password(form.password.data)
        db.session.add(admin)
        db.session.commit()
        flash(f'مدیر تایید "{admin.username}" با موفقیت ایجاد شد.', 'success')
        return redirect(url_for('admin_dashboard'))
    return render_template('admin/create_admin.html', form=form)

@app.route('/admin/services/create', methods=['GET', 'POST'])
@login_required
@system_manager_required
def create_service():
    """Create new service"""
    form = ServiceForm()
    if form.validate_on_submit():
        # Verify Google Doc access
        if google_docs_service:
            try:
                if not google_docs_service.verify_document_access(form.google_doc_id.data):
                    flash('خطا: دسترسی به Google Doc امکان‌پذیر نیست. لطفاً شناسه را بررسی کنید.', 'danger')
                    return render_template('admin/create_service.html', form=form)
            except Exception as e:
                flash(f'خطا در دسترسی به Google Doc: {str(e)}', 'danger')
                return render_template('admin/create_service.html', form=form)
        else:
            flash('خطا: سرویس Google Docs پیکربندی نشده است.', 'danger')
            return render_template('admin/create_service.html', form=form)
        
        service = Service(
            name=form.name.data,
            description=form.description.data,
            google_doc_id=form.google_doc_id.data,
            created_by=current_user.id
        )
        
        db.session.add(service)
        db.session.commit()
        flash('خدمت جدید با موفقیت ایجاد شد.', 'success')
        return redirect(url_for('edit_service_fields', service_id=service.id))
    return render_template('admin/create_service.html', form=form)

@app.route('/admin/services/preview-google-doc/<doc_id>')
@login_required
@system_manager_required
def preview_google_doc(doc_id):
    """Preview Google Doc content via AJAX"""
    if not google_docs_service:
        return jsonify({'error': 'سرویس Google Docs پیکربندی نشده است'}), 500
    
    try:
        # Extract content and placeholders
        doc_info = google_docs_service.extract_text_and_placeholders(doc_id)
        
        return jsonify({
            'success': True,
            'title': doc_info['title'],
            'paragraphs': doc_info['text_content'][:10],  # First 10 paragraphs
            'placeholders': doc_info['placeholders'],
            'total_paragraphs': len(doc_info['text_content']),
            'doc_url': google_docs_service.get_document_url(doc_id)
        })
        
    except Exception as e:
        app.logger.error(f"Error previewing Google Doc: {str(e)}")
        return jsonify({'error': f'خطا در خواندن سند: {str(e)}'}), 500



@app.route('/admin/services/<int:service_id>/edit', methods=['GET', 'POST'])
@login_required
@system_manager_required
def edit_service(service_id):
    """Edit service details"""
    service = Service.query.get_or_404(service_id)
    form = ServiceForm(obj=service)
    
    if form.validate_on_submit():
        # Verify Google Doc access if changed
        if form.google_doc_id.data != service.google_doc_id:
            if google_docs_service:
                try:
                    if not google_docs_service.verify_document_access(form.google_doc_id.data):
                        flash('خطا: دسترسی به Google Doc امکان‌پذیر نیست. لطفاً شناسه را بررسی کنید.', 'danger')
                        return render_template('admin/edit_service.html', form=form, service=service)
                except Exception as e:
                    flash(f'خطا در دسترسی به Google Doc: {str(e)}', 'danger')
                    return render_template('admin/edit_service.html', form=form, service=service)
            else:
                flash('خطا: سرویس Google Docs پیکربندی نشده است.', 'danger')
                return render_template('admin/edit_service.html', form=form, service=service)
        
        service.name = form.name.data
        service.description = form.description.data
        service.google_doc_id = form.google_doc_id.data
        
        db.session.commit()
        flash('خدمت با موفقیت به‌روزرسانی شد.', 'success')
        return redirect(url_for('admin_dashboard'))
    
    return render_template('admin/edit_service.html', form=form, service=service)

@app.route('/admin/services/<int:service_id>/fields', methods=['GET', 'POST'])
@login_required
@system_manager_required
def edit_service_fields(service_id):
    """Edit service form fields"""
    service = Service.query.get_or_404(service_id)
    form = FormFieldForm()
    
    if form.validate_on_submit():
        field = FormField(
            service_id=service.id,
            field_name=form.field_name.data,
            field_label=form.field_label.data,
            field_type=form.field_type.data,
            is_required=form.is_required.data,
            placeholder=form.placeholder.data,
            document_placeholder=form.document_placeholder.data,
            field_order=service.form_fields.count()
        )
        
        if form.field_type.data == 'select' and form.options.data:
            options = [opt.strip() for opt in form.options.data.split('\n') if opt.strip()]
            field.options = json.dumps(options, ensure_ascii=False)
        
        db.session.add(field)
        db.session.commit()
        flash('فیلد جدید اضافه شد.', 'success')
        return redirect(url_for('edit_service_fields', service_id=service.id))
    
    fields = service.form_fields.order_by(FormField.field_order).all()
    return render_template('admin/edit_fields.html', service=service, form=form, fields=fields)

@app.route('/admin/services/<int:service_id>/stats')
@login_required
@system_manager_required
def service_stats(service_id):
    """View detailed statistics for a service"""
    service = Service.query.get_or_404(service_id)
    stats = service.get_stats()
    
    # Get all requests for this service
    page = request.args.get('page', 1, type=int)
    requests = ServiceRequest.query.filter_by(service_id=service_id).order_by(
        ServiceRequest.created_at.desc()
    ).paginate(page=page, per_page=app.config['REQUESTS_PER_PAGE'])
    
    return render_template('admin/service_stats.html', 
                         service=service, 
                         stats=stats, 
                         requests=requests)

@app.route('/admin/services/<int:service_id>/detect-placeholders')
@login_required
@system_manager_required
def detect_placeholders(service_id):
    """Detect placeholders in the service's template"""
    service = Service.query.get_or_404(service_id)
    
    if not service.google_doc_id:
        return jsonify({'error': 'No template configured for this service'}), 400
    
    if not google_docs_service:
        return jsonify({'error': 'Google Docs service not configured'}), 500
    
    try:
        # Export template as DOCX to detect placeholders
        docx_data = google_docs_service.export_as_docx(service.google_doc_id)
        
        # Create a temporary file to analyze
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(docx_data)
            tmp_file_path = tmp_file.name
        
        try:
            from docx import Document
            doc = Document(tmp_file_path)
            
            # Detect placeholders
            placeholders = set()
            
            # Check paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text
                found = re.findall(r'\{\{([^}]+)\}\}', text)
                placeholders.update(found)
            
            # Check tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text
                            found = re.findall(r'\{\{([^}]+)\}\}', text)
                            placeholders.update(found)
            
            # Check headers and footers
            for section in doc.sections:
                for paragraph in section.header.paragraphs:
                    text = paragraph.text
                    found = re.findall(r'\{\{([^}]+)\}\}', text)
                    placeholders.update(found)
                for paragraph in section.footer.paragraphs:
                    text = paragraph.text
                    found = re.findall(r'\{\{([^}]+)\}\}', text)
                    placeholders.update(found)
            
            # Get existing field mappings
            existing_mappings = {}
            for field in service.form_fields:
                if field.document_placeholder:
                    existing_mappings[field.document_placeholder] = field.field_name
            
            return jsonify({
                'success': True,
                'placeholders': sorted(list(placeholders)),
                'existing_mappings': existing_mappings
            })
            
        finally:
            # Clean up temp file
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)
                
    except Exception as e:
        app.logger.error(f"Error detecting placeholders: {str(e)}")
        return jsonify({'error': f'Error detecting placeholders: {str(e)}'}), 500


# Approval Admin Routes
@app.route('/approver')
@login_required
@approval_admin_required
def approver_dashboard():
    """Approval admin dashboard"""
    page = request.args.get('page', 1, type=int)
    requests = ServiceRequest.query.filter_by(status='pending').order_by(
        ServiceRequest.created_at.desc()
    ).paginate(page=page, per_page=app.config['REQUESTS_PER_PAGE'])
    
    return render_template('approver/dashboard.html', 
                         requests=requests)

@app.route('/approver/request/<int:request_id>', methods=['GET', 'POST'])
@login_required
@approval_admin_required
def review_request(request_id):
    """Review and approve/reject a request"""
    service_request = ServiceRequest.query.get_or_404(request_id)
    form = ApprovalForm()
    
    if form.validate_on_submit():
        service_request.status = form.action.data + 'd'  # approved or rejected
        service_request.approval_note = form.note.data
        service_request.approved_by = current_user.id
        
        if form.action.data == 'approve':
            # Generate PDF using the template
            try:
                pdf_filename = generate_pdf_from_request(service_request)
                service_request.pdf_filename = pdf_filename
                flash('درخواست تایید شد و PDF تولید شد.', 'success')
            except Exception as e:
                # If PDF generation fails, don't approve the request
                service_request.status = 'pending'
                service_request.approval_note = None
                service_request.approved_by = None
                flash(f'خطا در تولید PDF: {str(e)}. درخواست تایید نشد.', 'danger')
                return redirect(url_for('review_request', request_id=request_id))
        else:
            flash('درخواست رد شد.', 'info')
        
        db.session.commit()
        return redirect(url_for('approver_dashboard'))
    
    form_data = service_request.get_form_data()
    return render_template('approver/review_request.html', 
                         request=service_request, 
                         form=form,
                         form_data=form_data)

# Public User Routes
@app.route('/service/<int:service_id>/request', methods=['GET', 'POST'])
def request_service(service_id):
    """Submit a service request"""
    service = Service.query.get_or_404(service_id)
    if not service.is_active:
        flash('این خدمت در حال حاضر غیرفعال است.', 'warning')
        return redirect(url_for('index'))
    
    # Create dynamic form
    class DynamicForm(ServiceRequestForm):
        pass
    
    # Add fields dynamically
    for field in service.form_fields.order_by(FormField.field_order):
        field_class = None
        validators = []
        
        if field.is_required:
            validators.append(DataRequired(message='این فیلد اجباری است'))
        
        if field.field_type == 'text':
            field_class = StringField
        elif field.field_type == 'number':
            field_class = IntegerField
        elif field.field_type == 'email':
            field_class = StringField
            validators.append(Email(message='ایمیل معتبر وارد کنید'))
        elif field.field_type == 'textarea':
            field_class = TextAreaField
        elif field.field_type == 'select':
            choices = [(opt, opt) for opt in field.get_options()]
            setattr(DynamicForm, field.field_name, 
                   SelectField(field.field_label, choices=choices, validators=validators))
            continue
        else:
            field_class = StringField
        
        if field_class:
            setattr(DynamicForm, field.field_name, 
                   field_class(field.field_label, validators=validators, 
                             render_kw={'placeholder': field.placeholder}))
    
    form = DynamicForm()
    
    if form.validate_on_submit():
        # Collect form data
        form_data = {}
        for field in service.form_fields:
            if hasattr(form, field.field_name):
                form_data[field.field_name] = getattr(form, field.field_name).data
        
        # Create request
        service_request = ServiceRequest(
            service_id=service.id,
            tracking_code=generate_tracking_code()
        )
        service_request.set_form_data(form_data)
        
        db.session.add(service_request)
        db.session.commit()
        
        flash(f'درخواست شما با کد پیگیری {service_request.tracking_code} ثبت شد.', 'success')
        return redirect(url_for('track_request', tracking_code=service_request.tracking_code))
    
    return render_template('user/request_service.html', service=service, form=form)

@app.route('/track', methods=['GET', 'POST'])
def track_request():
    """Track request status"""
    form = TrackingForm()
    tracking_code = request.args.get('tracking_code')
    
    if tracking_code:
        form.tracking_code.data = tracking_code
    
    if form.validate_on_submit() or tracking_code:
        code = form.tracking_code.data if form.validate_on_submit() else tracking_code
        service_request = ServiceRequest.query.filter_by(tracking_code=code).first()
        
        if service_request:
            return render_template('user/request_status.html', 
                                 request=service_request, 
                                 form_data=service_request.get_form_data())
        else:
            flash('کد پیگیری یافت نشد.', 'warning')
    
    return render_template('user/track.html', form=form)

@app.route('/download/<tracking_code>')
def download_pdf(tracking_code):
    """Download approved request PDF"""
    service_request = ServiceRequest.query.filter_by(tracking_code=tracking_code).first_or_404()
    
    if service_request.status != 'approved' or not service_request.pdf_filename:
        flash('فایل PDF موجود نیست.', 'warning')
        return redirect(url_for('track_request', tracking_code=tracking_code))
    
    pdf_path = os.path.join(app.config['PDF_OUTPUT_FOLDER'], service_request.pdf_filename)
    if os.path.exists(pdf_path):
        return send_file(pdf_path, as_attachment=True, 
                        download_name=f'request_{tracking_code}.pdf')
    else:
        flash('فایل PDF یافت نشد.', 'danger')
        return redirect(url_for('track_request', tracking_code=tracking_code))

# PDF Generation
def generate_pdf_from_request(service_request):
    """Generate PDF from approved request without modifying the original template"""
    service = service_request.service
    form_data = service_request.get_form_data()
    
    if not google_docs_service:
        raise Exception("Google Docs service not configured")
    
    # Use the service's Google Doc template
    template_id = service.google_doc_id
    if not template_id:
        raise Exception("No template configured for this service")
    
    try:
        # Initialize document processor
        doc_processor = DocumentProcessor()
        
        # Prepare replacements mapping
        replacements = {}
        for field in service.form_fields:
            if field.document_placeholder:
                value = str(form_data.get(field.field_name, ''))
                replacements[field.document_placeholder] = value
        
        # Try DOCX format first (better for preserving formatting)
        try:
            # Export template as DOCX
            docx_data = google_docs_service.export_as_docx(template_id)
            
            # Process the DOCX template with replacements
            filled_docx = doc_processor.process_docx_template(docx_data, replacements)
            
            # Generate PDF from the filled DOCX
            pdf_filename = f"output_{service_request.tracking_code}.pdf"
            pdf_path = os.path.join(app.config['PDF_OUTPUT_FOLDER'], pdf_filename)
            
            if doc_processor.generate_pdf_from_docx(filled_docx, pdf_path):
                return pdf_filename
            else:
                raise Exception("Failed to generate PDF from DOCX")
                
        except Exception as docx_error:
            app.logger.warning(f"DOCX processing failed, trying HTML: {str(docx_error)}")
            
            # Fallback to HTML format
            try:
                # Export template as HTML
                html_data = google_docs_service.export_as_html(template_id)
                
                # Process the HTML template with replacements
                filled_html = doc_processor.process_html_template(html_data, replacements)
                
                # Generate PDF from the filled HTML
                pdf_filename = f"output_{service_request.tracking_code}.pdf"
                pdf_path = os.path.join(app.config['PDF_OUTPUT_FOLDER'], pdf_filename)
                
                if doc_processor.generate_pdf_from_html(filled_html, pdf_path):
                    return pdf_filename
                else:
                    raise Exception("Failed to generate PDF from HTML")
                    
            except Exception as html_error:
                app.logger.error(f"HTML processing also failed: {str(html_error)}")
                
                # Final fallback: use Google Docs API directly to export as PDF
                # This won't have placeholder replacements but at least generates something
                pdf_data = google_docs_service.export_as_pdf(template_id)
                
                # Save PDF locally
                pdf_filename = f"output_{service_request.tracking_code}.pdf"
                pdf_path = os.path.join(app.config['PDF_OUTPUT_FOLDER'], pdf_filename)
                
                with open(pdf_path, 'wb') as f:
                    f.write(pdf_data)
                
                app.logger.warning("Generated PDF without placeholder replacements as fallback")
                return pdf_filename
        
    except Exception as e:
        app.logger.error(f"Error generating PDF: {str(e)}")
        raise
    


# Initialize database
@app.cli.command()
def init_db():
    """Initialize the database."""
    db.create_all()
    
    # Create default system manager if not exists
    if not User.query.filter_by(role='system_manager').first():
        admin = User(
            username='admin',
            email='admin@example.com',
            role='system_manager'
        )
        admin.set_password('admin123')
        db.session.add(admin)
        db.session.commit()
        print("Default system manager created: username='admin', password='admin123'")
    
    print("Database initialized!")



# Error handlers
@app.errorhandler(404)
def not_found_error(error):
    return render_template('errors/404.html'), 404

@app.errorhandler(500)
def internal_error(error):
    db.session.rollback()
    return render_template('errors/500.html'), 500

if __name__ == '__main__':
    with app.app_context():
        db.create_all()
        
        # Create default system manager if not exists
        if not User.query.filter_by(role='system_manager').first():
            admin = User(
                username='admin',
                email='admin@example.com',
                role='system_manager'
            )
            admin.set_password('admin123')
            db.session.add(admin)
            db.session.commit()
            print("Default system manager created: username='admin', password='admin123'")
    
    app.run(debug=True)