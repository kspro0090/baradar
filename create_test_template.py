#!/usr/bin/env python3
"""
Create a test Word template with placeholders for testing
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import os

def create_test_template():
    """Create a test template with various placeholders"""
    
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading('نامه درخواست خدمت', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add date placeholder
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_para.add_run('تاریخ: {{date}}')
    
    # Add greeting
    doc.add_paragraph()
    greeting = doc.add_paragraph('با سلام و احترام')
    greeting.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add main content with placeholders
    main_para = doc.add_paragraph()
    main_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    main_para.add_run('اینجانب ')
    name_run = main_para.add_run('{{employee_name}}')
    name_run.bold = True
    main_para.add_run(' با کد پرسنلی ')
    code_run = main_para.add_run('{{employee_code}}')
    code_run.bold = True
    main_para.add_run(' از بخش ')
    dept_run = main_para.add_run('{{department}}')
    dept_run.bold = True
    main_para.add_run(' درخواست ')
    service_run = main_para.add_run('{{service_type}}')
    service_run.bold = True
    main_para.add_run(' را دارم.')
    
    # Add details section
    doc.add_paragraph()
    details_heading = doc.add_heading('جزئیات درخواست', level=2)
    details_heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add a table with placeholders
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[1].text = 'عنوان'
    hdr_cells[0].text = 'مقدار'
    
    # Table data with placeholders
    table_data = [
        ('نوع درخواست', '{{request_type}}'),
        ('تاریخ شروع', '{{start_date}}'),
        ('تاریخ پایان', '{{end_date}}'),
        ('توضیحات', '{{description}}')
    ]
    
    for i, (label, placeholder) in enumerate(table_data, 1):
        row_cells = table.rows[i].cells
        row_cells[1].text = label
        row_cells[0].text = placeholder
    
    # Set table alignment
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add contact info
    doc.add_paragraph()
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact_para.add_run('ایمیل: ')
    contact_para.add_run('{{email}}')
    contact_para.add_run(' | ')
    contact_para.add_run('تلفن: ')
    contact_para.add_run('{{phone}}')
    
    # Add signature section
    doc.add_paragraph()
    doc.add_paragraph()
    signature = doc.add_paragraph('امضا: {{employee_name}}')
    signature.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Save the document
    output_dir = 'test_templates'
    os.makedirs(output_dir, exist_ok=True)
    
    filename = os.path.join(output_dir, 'service_request_template.docx')
    doc.save(filename)
    
    print(f"Test template created: {filename}")
    print("\nPlaceholders used:")
    placeholders = [
        'date', 'employee_name', 'employee_code', 'department',
        'service_type', 'request_type', 'start_date', 'end_date',
        'description', 'email', 'phone'
    ]
    for p in placeholders:
        print(f"  - {{{{{p}}}}}")
    
    return filename

if __name__ == '__main__':
    create_test_template()